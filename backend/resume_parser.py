"""
Resume Parser Module
Handles extraction of text from PDF and DOCX files.
Includes robust error handling and text cleaning.
"""

import re
from pathlib import Path
from typing import Optional, Tuple
import logging

# PDF parsing
try:
    import pdfplumber
    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    PDF_PLUMBER_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    A robust resume parser that extracts text from PDF and DOCX files.
    Uses multiple fallback methods for reliability.
    """
    
    def __init__(self):
        """Initialize the parser and check available libraries."""
        self.supported_formats = []
        
        if PDF_PLUMBER_AVAILABLE or PYPDF2_AVAILABLE:
            self.supported_formats.extend(['.pdf'])
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx', '.doc'])
        self.supported_formats.append('.txt')
        
        logger.info(f"ResumeParser initialized. Supported formats: {self.supported_formats}")
    
    def extract_text(self, file_path: str) -> Tuple[str, bool]:
        """
        Extract text from a resume file.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Tuple of (extracted_text, success_flag)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return "", False
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                logger.error(f"Unsupported file format: {extension}")
                return "", False
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return "", False
    
    def _extract_from_pdf(self, file_path: Path) -> Tuple[str, bool]:
        """
        Extract text from PDF using multiple methods.
        Primary: pdfplumber, Fallback: PyPDF2
        """
        text = ""
        
        # Try pdfplumber first (better at handling complex layouts)
        if PDF_PLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(page_text)
                    text = "\n".join(pages_text)
                    
                if text.strip():
                    logger.info(f"Successfully extracted text using pdfplumber: {file_path.name}")
                    return self._clean_text(text), True
            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path.name}: {str(e)}")
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                reader = PdfReader(file_path)
                pages_text = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
                
                if text.strip():
                    logger.info(f"Successfully extracted text using PyPDF2: {file_path.name}")
                    return self._clean_text(text), True
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {file_path.name}: {str(e)}")
        
        logger.error(f"All PDF extraction methods failed for {file_path.name}")
        return "", False
    
    def _extract_from_docx(self, file_path: Path) -> Tuple[str, bool]:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return "", False
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            text = "\n".join(paragraphs)
            logger.info(f"Successfully extracted text from DOCX: {file_path.name}")
            return self._clean_text(text), True
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path.name}: {str(e)}")
            return "", False
    
    def _extract_from_txt(self, file_path: Path) -> Tuple[str, bool]:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Successfully read TXT file with {encoding} encoding: {file_path.name}")
                    return self._clean_text(text), True
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode TXT file with any encoding: {file_path.name}")
            return "", False
            
        except Exception as e:
            logger.error(f"TXT extraction failed for {file_path.name}: {str(e)}")
            return "", False
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing noise and normalizing whitespace.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Replace multiple whitespace with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\-\.\,\@\+\#\/\(\)\:\;\&]', '', text)
        
        # Remove multiple periods or dashes
        text = re.sub(r'\.{2,}', '.', text)
        text = re.sub(r'\-{2,}', '-', text)
        
        # Normalize whitespace around punctuation
        text = re.sub(r'\s+([,.\-])', r'\1', text)
        text = re.sub(r'([,.\-])\s+', r'\1 ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_word_count(self, text: str) -> int:
        """Get word count of text."""
        if not text:
            return 0
        return len(text.split())
    
    def validate_resume_content(self, text: str) -> Tuple[bool, str]:
        """
        Validate that extracted text appears to be a valid resume.
        
        Args:
            text: Extracted text to validate
            
        Returns:
            Tuple of (is_valid, message)
        """
        if not text:
            return False, "No text content extracted"
        
        word_count = self.get_word_count(text)
        
        if word_count < 50:
            return False, f"Content too short ({word_count} words). Minimum 50 words expected."
        
        if word_count > 10000:
            return False, f"Content too long ({word_count} words). Maximum 10000 words expected."
        
        # Check for common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'work', 'project',
            'email', 'phone', 'address', 'summary', 'objective',
            'certification', 'achievement', 'responsibility'
        ]
        
        text_lower = text.lower()
        matches = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        if matches < 2:
            return False, "Content does not appear to be a resume (missing common sections)"
        
        return True, "Valid resume content"


# Global parser instance
resume_parser = ResumeParser()